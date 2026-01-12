import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt


st.set_page_config(
    page_title="PDL Setlist",
    layout="wide"
)


def gerar_docx(df, col_musica, col_artista, col_bpm, col_tom, col_obs):
    doc = Document()
    primeira = True

    for _, row in df.iterrows():
        titulo = str(row[col_musica]) if col_musica else ""
        if not titulo or titulo == "nan":
            continue  # pula linhas vazias

        if not primeira:
            doc.add_page_break()
        primeira = False

        artista = str(row[col_artista]) if col_artista else ""
        bpm = row[col_bpm] if col_bpm else ""
        tom = str(row[col_tom]) if col_tom else ""
        obs = str(row[col_obs]) if col_obs else ""

        # Normalizar valores "nan"
        artista = "" if artista.lower() == "nan" else artista
        tom = "" if tom.lower() == "nan" else tom
        obs = "" if obs.lower() == "nan" else obs

        # ---- BPM (linha 1) ----
        if bpm != "" and str(bpm).lower() != "nan":
            p_bpm_label = doc.add_paragraph()
            run_label = p_bpm_label.add_run("BPM")
            run_label.bold = True

            p_bpm_val = doc.add_paragraph()
            run_val = p_bpm_val.add_run(str(int(bpm)) if isinstance(bpm, (int, float)) else str(bpm))
            run_val.font.size = Pt(10)

        # ---- Observações (linha 2) ----
        if obs:
            p_obs = doc.add_paragraph()
            run_obs = p_obs.add_run(obs)
            run_obs.font.size = Pt(10)

        # ---- Título da música (grande) ----
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(titulo)
        run_title.bold = True
        run_title.font.size = Pt(18)

        # ---- Artista (linha abaixo, fonte 10) ----
        if artista:
            p_artist = doc.add_paragraph()
            run_artist = p_artist.add_run(artista)
            run_artist.font.size = Pt(10)

        # ---- Tom ----
        if tom:
            p_key = doc.add_paragraph()
            run_key_label = p_key.add_run("Tom: ")
            run_key_label.bold = True
            run_key_val = p_key.add_run(tom)
            run_key_val.font.size = Pt(10)

        # Espaço para onde você vai colocar a cifra depois, se quiser
        doc.add_paragraph()

    # Salva em memória para download
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def main():
    st.title("PDL Setlist – Gerador de folhas em .docx")

    st.markdown(
        """
        ### Como usar
        1. Exporte sua **setlist em CSV** (por exemplo do Excel).
        2. Faça o upload aqui.
        3. Escolha quais colunas são **Música, Artista, BPM, Tom e Observações**.
        4. Clique em **Gerar .docx** e baixe o arquivo para imprimir.

        > Dica: deixe uma coluna com a ordem das músicas no seu CSV, se quiser manter a sequência.
        """
    )

    uploaded_file = st.file_uploader("Envie o arquivo de setlist (.csv)", type=["csv"])

    if uploaded_file is not None:
        try:
            df = pd.read_csv(uploaded_file)
        except Exception:
            uploaded_file.seek(0)
            df = pd.read_csv(uploaded_file, sep=";")

        st.subheader("Pré-visualização do CSV")
        st.dataframe(df, use_container_width=True)

        st.markdown("#### Selecione as colunas")

        cols = ["(nenhuma)"] + list(df.columns)

        col1, col2, col3 = st.columns(3)
        with col1:
            col_musica = st.selectbox("Coluna da MÚSICA", cols, index=cols.index(cols[1]) if len(cols) > 1 else 0)
            col_artista = st.selectbox("Coluna do ARTISTA", cols)
        with col2:
            col_bpm = st.selectbox("Coluna do BPM", cols)
            col_tom = st.selectbox("Coluna do TOM", cols)
        with col3:
            col_obs = st.selectbox("Coluna das OBSERVAÇÕES", cols)

        # Converter "(nenhuma)" para None
        col_musica = None if col_musica == "(nenhuma)" else col_musica
        col_artista = None if col_artista == "(nenhuma)" else col_artista
        col_bpm = None if col_bpm == "(nenhuma)" else col_bpm
        col_tom = None if col_tom == "(nenhuma)" else col_tom
        col_obs = None if col_obs == "(nenhuma)" else col_obs

        if not col_musica:
            st.error("Você precisa selecionar pelo menos a coluna da **MÚSICA**.")
            return

        if st.button("Gerar arquivo .docx"):
            buffer = gerar_docx(df, col_musica, col_artista, col_bpm, col_tom, col_obs)

            st.success("Arquivo gerado com sucesso! Clique abaixo para baixar.")
            st.download_button(
                label="Baixar setlist.docx",
                data=buffer,
                file_name="PDL-setlist.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


if __name__ == "__main__":
    main()
